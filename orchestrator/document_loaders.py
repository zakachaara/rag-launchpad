import os
from pathlib import Path
from typing import List
from dataclasses import dataclass, field

from logger import get_logger

log = get_logger("loaders")


@dataclass
class Document:
    page_content: str
    metadata: dict = field(default_factory=dict)


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def load_txt(path: str) -> Document:
    log.debug("Reading TXT  %s", path)
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    log.debug("  → %d chars", len(text))
    return Document(page_content=text, metadata={"source": path})


def load_pdf(path: str) -> List[Document]:
    log.debug("Reading PDF  %s", path)
    try:
        from pypdf import PdfReader
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore

    reader = PdfReader(path)
    total_pages = len(reader.pages)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": path, "page": i + 1},
            ))
    log.debug("  → %d/%d pages had extractable text", len(docs), total_pages)
    return docs


def load_docx(path: str) -> Document:
    log.debug("Reading DOCX  %s", path)
    from docx import Document as DocxDocument  # type: ignore

    doc = DocxDocument(path)
    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    log.debug("  → %d chars from %d paragraphs", len(text), len(doc.paragraphs))
    return Document(page_content=text, metadata={"source": path})


def load_documents_from_folder(folder: str) -> List[Document]:
    """Recursively scan folder for supported files and return Document objects."""
    folder_path = Path(folder)
    if not folder_path.exists():
        raise FileNotFoundError(f"Data folder not found: {folder}")

    all_files = [
        p for p in folder_path.rglob("*")
        if p.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    log.info(
        "Scanning '%s' — found %d supported file(s): %s",
        folder,
        len(all_files),
        [p.name for p in all_files],
    )

    documents: List[Document] = []
    skipped = 0

    for file_path in all_files:
        path_str = str(file_path)
        try:
            ext = file_path.suffix.lower()
            if ext == ".txt":
                documents.append(load_txt(path_str))
            elif ext == ".pdf":
                documents.extend(load_pdf(path_str))
            elif ext == ".docx":
                documents.append(load_docx(path_str))
        except Exception as e:
            skipped += 1
            log.warning("Skipping unreadable file '%s': %s", path_str, e)

    log.info(
        "Document loading complete — %d doc(s) loaded, %d skipped",
        len(documents), skipped,
    )
    return documents
